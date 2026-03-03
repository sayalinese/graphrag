import os
import logging
import threading
import subprocess
from typing import List, Optional, Dict, Any
from pathlib import Path

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyMuPDFLoader, 
    Docx2txtLoader, 
    UnstructuredExcelLoader,
    TextLoader
)

from app import db
from app.models.knowledge_base import Document as DocumentModel, KnowledgeBase
from app.services.rag.kb_embedding import KBEmbeddingService
from .splitting import smart_split_md_html, smart_split_docx, smart_split_pdf, smart_split_excel, semantic_refine_with_embeddings

logger = logging.getLogger(__name__)

COM_WORD_LOCK = threading.Lock()  # 避免并发多次唤起 Word COM 造成崩溃


class DocumentManager:
    """文档管理器 - 负责文档的增删查改（pgvector 版）"""
    
    def __init__(self, embedding_service: Optional[KBEmbeddingService] = None):
        """
        初始化文档管理器
        
        Args:
            embedding_service: embedding服务实例，如果为None则创建默认实例
        """
        self.embedding_service = embedding_service or KBEmbeddingService()
    
    def convert_doc_to_docx(self, doc_path: str) -> str:
        """将 .doc 转为 .docx，提供多层回退：
        1) win32com + Word（需要本地安装 Word，线程内 CoInitialize）
        2) LibreOffice (soffice) 命令行转换
        3) （放弃）直接抛错，提示用户手动另存为 docx
        注：python-docx 不支持旧 .doc，故不再尝试其读取 .doc。
        """
        from werkzeug.utils import secure_filename
        file_dir = os.path.dirname(doc_path)
        file_name = os.path.splitext(os.path.basename(doc_path))[0]
        safe_name = secure_filename(file_name) or "doc_file"
        safe_doc_path = os.path.join(file_dir, f"{safe_name}.doc")
        if safe_doc_path != doc_path:
            try:
                import shutil
                shutil.copyfile(doc_path, safe_doc_path)
            except Exception as _copy_e:
                logger.warning(f"Copy .doc temp failed, use original: { _copy_e }")
                safe_doc_path = doc_path
        docx_path = os.path.join(file_dir, f"{safe_name}.docx")

        # 已存在直接返回
        if os.path.exists(docx_path):
            return docx_path

        #  win32com
        com_tried_error: Optional[str] = None
        try:
            import pythoncom  # type: ignore
            import win32com.client  # type: ignore
            with COM_WORD_LOCK:
                pythoncom.CoInitialize()
                word = None
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(safe_doc_path)
                    doc.SaveAs2(docx_path, FileFormat=16)
                    doc.Close()
                    logger.info(f"[DOC->DOCX] win32com converted {safe_doc_path} -> {docx_path}")
                    return docx_path
                finally:
                    try:
                        if word:
                            word.Quit()
                    except Exception:
                        pass
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass
        except ImportError:
            com_tried_error = "win32com/pythoncom not installed"
        except Exception as e:
            com_tried_error = f"win32com failed: {e}"
            logger.warning(f"[DOC->DOCX] win32com path failed: {e}")

        # LibreOffice (soffice)
        soffice_error: Optional[str] = None
        try:
            # Windows 下一般在 PATH 手动加入 LibreOffice program 目录
            result = subprocess.run([
                'soffice', '--headless', '--convert-to', 'docx', '--outdir', file_dir, safe_doc_path
            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=60)
            if result.returncode == 0 and os.path.exists(docx_path):
                logger.info(f"[DOC->DOCX] soffice converted {safe_doc_path} -> {docx_path}")
                return docx_path
            else:
                soffice_error = f"code={result.returncode} stdout={result.stdout.strip()} stderr={result.stderr.strip()}"
        except FileNotFoundError:
            soffice_error = "soffice not found"
        except subprocess.TimeoutExpired:
            soffice_error = "soffice timeout"
        except Exception as e:
            soffice_error = f"soffice error: {e}"

        # 清理临时安全 doc
        if safe_doc_path != doc_path and os.path.exists(safe_doc_path):
            try:
                os.remove(safe_doc_path)
            except Exception:
                pass

        # 全部失败
        detail_parts = []
        if com_tried_error:
            detail_parts.append(com_tried_error)
        if soffice_error:
            detail_parts.append(soffice_error)
        detail = "; ".join(detail_parts) or "no method attempted"
        logger.error(f"Failed to convert .doc file {doc_path}: {detail}")
        raise ValueError(
            "Failed to convert .doc file. 请安装 Microsoft Word (win32com) 或 LibreOffice 并加入 PATH，或手动另存为 .docx 后重新上传. Details: " + detail
        )

    def load_document(self, file_path: str) -> List[Document]:
        """
        根据文件类型加载文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档列表
        """
        file_path = Path(file_path)
        
        try:
            if file_path.suffix.lower() == '.pdf':
                loader = PyMuPDFLoader(str(file_path))
                return loader.load()
            elif file_path.suffix.lower() == '.doc':
                # 先将.doc转换为.docx
                logger.info(f"Converting .doc file to .docx: {file_path}")
                docx_path = self.convert_doc_to_docx(str(file_path))
                # 加载转换后的 .docx，并将元数据的 source 指回原 .doc，另存 processed_source 便于后续 .docx 结构化切分
                docs = self.load_document(docx_path)
                try:
                    import os as _os
                    original_src = str(file_path)
                    processed_src = str(docx_path)
                    for d in docs:
                        md = dict(getattr(d, 'metadata', {}) or {})
                        md['source'] = original_src
                        md['processed_source'] = processed_src
                        md['filename'] = _os.path.basename(original_src)
                        d.metadata = md
                except Exception:
                    pass
                return docs
            elif file_path.suffix.lower() == '.docx':
                # 尝试多种方法加载Word文档
                errors = []
                
                # 方法1: Docx2txtLoader
                try:
                    loader = Docx2txtLoader(str(file_path))
                    docs = loader.load()
                    # 统一补充/覆盖 source 元数据，确保为 .docx 实际路径
                    for d in docs:
                        try:
                            md = dict(getattr(d, 'metadata', {}) or {})
                            md['source'] = str(file_path)
                            d.metadata = md
                        except Exception:
                            pass
                    if docs and docs[0].page_content.strip():
                        return docs
                except Exception as e:
                    errors.append(f"Docx2txtLoader: {str(e)}")
                
                # 方法2: python-docx
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(str(file_path))
                    text = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text.strip() + "\n"
                    if text.strip():
                        return [Document(page_content=text, metadata={"source": str(file_path)})]
                except Exception as e:
                    errors.append(f"python-docx: {str(e)}")
                
                # 方法3: 直接读取文本
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    if text.strip():
                        return [Document(page_content=text, metadata={"source": str(file_path)})]
                except Exception as e:
                    errors.append(f"Text reading: {str(e)}")
                
                # 所有方法都失败了
                raise ValueError(f"Failed to load Word document. Errors: {'; '.join(errors)}")
                
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                    # Excel 双写：sheet 级 markdown + 行级记录
                    try:
                        import pandas as pd  # type: ignore
                        xls = pd.read_excel(str(file_path), sheet_name=None, dtype=str)
                        docs: List[Document] = []
                        def _esc(val: str) -> str:
                            return str(val if val is not None else "").replace("|","\\|").replace("\r"," ").replace("\n"," ").strip()
                        for sheet_name, df in (xls or {}).items():
                            try:
                                if df is None:
                                    continue
                                df = df.fillna("")
                                cols_raw_all = list(df.columns)
                                cols_raw = [c for c in cols_raw_all if str(c).strip() and not str(c).strip().lower().startswith('unnamed') and str(c).strip().lower() != 'index']
                                if cols_raw:
                                    df = df[cols_raw]
                                cols = [_esc(c) for c in cols_raw]
                                # sheet 级 markdown
                                lines = []
                                if cols:
                                    header = "| " + " | ".join(cols) + " |"
                                    divider = "| " + " | ".join(["---" for _ in cols]) + " |"
                                    lines.append(header)
                                    lines.append(divider)
                                    for _, row in df.iterrows():
                                        vals = [_esc(row.get(c, "")) for c in cols_raw]
                                        lines.append("| " + " | ".join(vals) + " |")
                                else:
                                    lines.append(str(df))
                                sheet_content = "\n".join(lines).strip() + "\n"
                                if sheet_content:
                                    md_sheet = {
                                        "source": str(file_path),
                                        "sheet": str(sheet_name),
                                        "h2": str(sheet_name),
                                        "h3": " | ".join(cols) if cols else None,
                                        "start_index": 0,
                                        "chunk_type": "sheet"
                                    }
                                    docs.append(Document(page_content=sheet_content, metadata=md_sheet))
                                # 行级
                                headers_join = " | ".join(cols) if cols else None
                                for ridx, row in df.iterrows():
                                    row_vals = [_esc(row.get(c, "")) for c in cols_raw]
                                    row_line = " | ".join(row_vals)
                                    if not row_line.strip():
                                        continue
                                    md_row = {
                                        "source": str(file_path),
                                        "sheet": str(sheet_name),
                                        "h2": str(sheet_name),
                                        "h3": headers_join,
                                        "row_index": int(ridx),
                                        "start_index": int(ridx),
                                        "chunk_type": "row"
                                    }
                                    docs.append(Document(page_content=row_line, metadata=md_row))
                            except Exception:
                                continue
                        if docs:
                            return docs
                    except Exception:
                        pass
                    loader = UnstructuredExcelLoader(str(file_path))
                    loaded = loader.load()
                    for d in loaded:
                        try:
                            md = dict(getattr(d, 'metadata', {}) or {})
                            md['chunk_type'] = md.get('chunk_type') or 'sheet'
                            d.metadata = md
                        except Exception:
                            pass
                    return loaded
            elif file_path.suffix.lower() in ['.txt', '.md', '.html', '.htm']:
                loader = TextLoader(str(file_path), encoding='utf-8')
                return loader.load()
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    def process_document_for_kb(self, file_path: str, kb_id: int,
                                vector_store_path: Optional[str] = None,
                                semantic_options: Optional[Dict[str, Any]] = None,
                                skip_faiss: bool = True,
                                return_chunks_only: bool = False,
                                clean_invalid: bool = False) -> Dict[str, Any] | List[Document]:
        """
        为知识库处理文档（pgvector 版）
        
        Args:
            file_path: 文档文件路径
            kb_id: 知识库ID
            vector_store_path: 向量存储路径（pgvector 暂不使用本地路径）
            semantic_options: 语义选项
            skip_faiss: 已废弃
            return_chunks_only: 仅返回分块
            clean_invalid: 清理无效块
            
        Returns:
            处理结果
        """
        try:
            # 加载文档
            documents = self.load_document(file_path)
            logger.info(f"Loaded {len(documents)} pages from {file_path}")
            # 统一补充基础元数据：source & filename
            try:
                import os as _os
                for _d in documents:
                    md = dict(getattr(_d, 'metadata', {}) or {})
                    src = md.get('source') or str(file_path)
                    md['source'] = str(src)
                    md['filename'] = _os.path.basename(str(src))
                    _d.metadata = md
            except Exception:
                pass
            
            # 切分文档：对 pdf/md/html/word/excel 优先用标题层级智能切分
            try:
                # 1) PDF 智能切分（字体大小启发式）
                smart_chunks = smart_split_pdf(documents, self.embedding_service.text_splitter)
                # 2) Markdown/HTML 智能切分
                if not smart_chunks:
                    smart_chunks = smart_split_md_html(documents, self.embedding_service.text_splitter)
                # 3) Word 智能切分
                if not smart_chunks:
                    smart_chunks = smart_split_docx(documents, self.embedding_service.text_splitter)
                # 4) Excel 智能切分（表格行级切分，保留header上下文）
                if not smart_chunks:
                    smart_chunks = smart_split_excel(documents, self.embedding_service.text_splitter)
            except Exception:
                smart_chunks = []
            if smart_chunks:
                split_docs = smart_chunks
            else:
                split_docs = self.embedding_service.split_documents(documents)

            # 语义边界检测细化：对较长块进一步基于相邻句相似度切分
            sem_enabled = True
            sem_threshold = 0.6
            if semantic_options and isinstance(semantic_options, dict):
                sem_enabled = bool(semantic_options.get('enabled', True))
                try:
                    sem_threshold = float(semantic_options.get('threshold', 0.6))
                except Exception:
                    sem_threshold = 0.6
            if sem_enabled:
                try:
                    # 检查是否为Excel智能切分的结果，如果是则跳过语义分割
                    is_excel_split = any(
                        d.metadata and d.metadata.get('sheet') and d.metadata.get('h3')
                        for d in split_docs
                    )
                    if not is_excel_split:
                        split_docs = semantic_refine_with_embeddings(
                            split_docs,
                            self.embedding_service,
                            target_chunk_size=getattr(self.embedding_service.text_splitter, 'chunk_size', 600) or 600,
                            overlap=getattr(self.embedding_service.text_splitter, 'chunk_overlap', 120) or 120,
                            sim_threshold=sem_threshold,
                        )
                except Exception:
                    pass
            # 可选：过滤"无效"纯符号/噪声块（例如仅标点、过短、无字母或中文）
            original_chunk_count = len(split_docs)
            removed_invalid = 0
            if clean_invalid:
                try:
                    import re
                    def _is_invalid(txt: str) -> bool:
                        if not txt:
                            return True
                        t = txt.strip()
                        if not t:
                            return True
                        # 特例: 诸如 "- 1 -"、"- 2 -" 之类的页眉/页脚或装饰
                        if re.fullmatch(r"[-–—\s]*\d{1,3}[-–—\s]*", t):
                            return True
                        if re.fullmatch(r"[-–—\s]*[A-Za-z]?\d{1,2}[A-Za-z]?[-–—\s]*", t) and len(t) <= 6:
                            return True
                        if re.fullmatch(r"-\s*\d{1,3}\s*-", t):  # '- 1 -'
                            return True
                        # 规则1: 很短且没有字母/中文（纯符号或数字+符号）
                        if len(t) <= 10 and not re.search(r"[A-Za-z\u4e00-\u9fff]", t):
                            # 若全为数字+标点也删除
                            if re.fullmatch(r"[0-9\s\-–—_=*~·.,，。:;|/\\()\[\]{}<>]+", t):
                                return True
                        # 规则2: 仅由标点/分隔符/数字构成
                        if re.fullmatch(r"[0-9\s\-–—_=*~·.,，。:;|/\\()\[\]{}<>]+", t):
                            return True
                        # 提取核心（字母/数字/中文）
                        core = re.findall(r"[A-Za-z0-9\u4e00-\u9fff]", t)
                        if not core:
                            return True
                        # 规则3: 核心字符全是数字且整体很短（列表编号/噪声）
                        if len(t) <= 10 and all(c.isdigit() for c in core):
                            return True
                        # 规则4: 核心占比过低（结构噪声）
                        ratio = len(core)/len(t)
                        if ratio < 0.18 and len(t) < 60:
                            return True
                        return False
                    filtered = []
                    for d in split_docs:
                        if _is_invalid(d.page_content or ''):
                            removed_invalid += 1
                            continue
                        filtered.append(d)
                    split_docs = filtered
                except Exception:
                    pass
            logger.info(f"Split into {len(split_docs)} chunks (clean_invalid={clean_invalid} removed={removed_invalid} of {original_chunk_count})")

            # 为每个块生成向量并放入 metadata['embedding']（pgvector 存储）
            pgvector_indexed = False
            pgvector_error: str | None = None
            try:
                texts_for_embed = [d.page_content for d in split_docs]
                if texts_for_embed:
                    vectors = self.embedding_service.embeddings.embed_documents(texts_for_embed)
                    for d, v in zip(split_docs, vectors):
                        try:
                            md = dict(getattr(d, 'metadata', {}) or {})
                            md['embedding'] = v
                            d.metadata = md
                        except Exception:
                            pass
                    logger.info(f"Generated embeddings for {len(vectors)} chunks")
                    pgvector_indexed = True
                else:
                    pgvector_error = "no texts for embedding"
            except Exception as _e:
                pgvector_error = f"embedding generation failed: {_e}"
                logger.warning(f"[pgvector] {pgvector_error}")
            
            # 更新数据库中的文档状态
            filename = os.path.basename(file_path)
            doc_model = DocumentModel.query.filter_by(
                filename=filename, 
                kb_id=kb_id
            ).first()
            
            if doc_model:
                try:
                    doc_model.status = 'COMPLETED'
                    db.session.commit()
                    logger.info(f"Document {filename} status updated to COMPLETED")
                except Exception as _e:
                    # 尝试回滚并记录，但不阻断主流程
                    try:
                        db.session.rollback()
                    except Exception:
                        pass
                    logger.warning(f"Commit COMPLETED failed for {filename}: {_e}")
            
            if return_chunks_only:
                return split_docs
            else:
                return {
                    'success': True,
                    'document_count': len(documents),
                    'chunk_count': len(split_docs),
                    'removed_invalid': removed_invalid,
                    'original_chunk_count': original_chunk_count,
                    'pgvector_indexed': pgvector_indexed,
                    'pgvector_error': pgvector_error,
                    'faiss_used': False,
                    'preview_available': bool(pgvector_indexed)
                }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            
            # 更新文档状态为失败
            filename = os.path.basename(file_path)
            try:
                # 若当前事务处于回滚状态，先回滚
                try:
                    db.session.rollback()
                except Exception:
                    pass
                doc_model = DocumentModel.query.filter_by(
                    filename=filename, 
                    kb_id=kb_id
                ).first()
                if doc_model:
                    doc_model.status = 'FAILED'
                    db.session.commit()
            except Exception as _e2:
                try:
                    db.session.rollback()
                except Exception:
                    pass
                logger.warning(f"Mark FAILED also failed for {filename}: {_e2}")
            
            return {
                'success': False,
                'error': str(e)
            }

    def simulate_split(self,
                       file_path: str,
                       chunk_size: int = 512,
                       chunk_overlap: int = 120,
                       semantic_refine: bool = True,
                       semantic_threshold: float = 0.6,
                       clean_invalid: bool = False,
                       max_preview_chars: int = 50000,
                       max_chunks: int = 800) -> Dict[str, Any]:
        """对单个文件执行与正式入库近似的切分流程，仅返回结果用于前端预览。"""
        try:
            # 临时覆写 splitter（不修改实例全局配置，生成一个新的）
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            custom_splitter = RecursiveCharacterTextSplitter(
                chunk_size=max(50, int(chunk_size)),
                chunk_overlap=max(0, int(min(chunk_overlap, chunk_size-1))) if chunk_size > 1 else 0,
                length_function=len,
                add_start_index=True,
            )

            # 加载
            documents = self.load_document(file_path)
            if not documents:
                return {'success': False, 'error': 'empty document'}

            # 若是普通非结构化文本且太大，截断影响性能（结构化智能切分路径直接使用全部）
            plain_docs = []
            for d in documents:
                if (getattr(d, 'metadata', {}) or {}).get('sheet') or (getattr(d, 'metadata', {}) or {}).get('h1'):
                    plain_docs.append(d)
                else:
                    txt = (d.page_content or '')
                    if len(txt) > max_preview_chars:
                        d.page_content = txt[:max_preview_chars]
                    plain_docs.append(d)

            # 智能切分顺序复用，与正式流程保持一致（使用 custom_splitter）
            from .splitting import smart_split_pdf, smart_split_md_html, smart_split_docx, smart_split_excel, semantic_refine_with_embeddings
            try:
                smart_chunks = smart_split_pdf(plain_docs, custom_splitter)
                if not smart_chunks:
                    smart_chunks = smart_split_md_html(plain_docs, custom_splitter)
                if not smart_chunks:
                    smart_chunks = smart_split_docx(plain_docs, custom_splitter)
                if not smart_chunks:
                    smart_chunks = smart_split_excel(plain_docs, custom_splitter)
            except Exception:
                smart_chunks = []
            if smart_chunks:
                split_docs: List[Document] = smart_chunks
            else:
                # 回退递归切分
                try:
                    split_docs = custom_splitter.split_documents(plain_docs)
                except Exception:
                    split_docs = plain_docs

            # 语义细化（可选；Excel 行级跳过）
            if semantic_refine:
                try:
                    is_excel = any(d.metadata and d.metadata.get('sheet') and d.metadata.get('h3') for d in split_docs)
                    if not is_excel:
                        split_docs = semantic_refine_with_embeddings(
                            split_docs,
                            self.embedding_service,
                            target_chunk_size=chunk_size,
                            overlap=chunk_overlap,
                            sim_threshold=float(semantic_threshold),
                        )
                except Exception:
                    pass

            # 可选清洗
            removed_invalid = 0
            original_chunk_count = len(split_docs)
            if clean_invalid:
                try:
                    import re
                    def _is_invalid(txt: str) -> bool:
                        if not txt:
                            return True
                        t = txt.strip()
                        if not t:
                            return True
                        # 页眉/页脚或装饰线，如 "- 1 -"、"-- 12 --"、" 3 "
                        if re.fullmatch(r"[-–—\s]*\d{1,3}[-–—\s]*", t):
                            return True
                        if re.fullmatch(r"[-–—\s]*[A-Za-z]?\d{1,2}[A-Za-z]?[-–—\s]*", t) and len(t) <= 6:
                            return True
                        if re.fullmatch(r"-\s*\d{1,3}\s*-", t):  # '- 1 -'
                            return True
                        # 很短且无字母/中文
                        if len(t) <= 10 and not re.search(r"[A-Za-z\u4e00-\u9fff]", t):
                            if re.fullmatch(r"[0-9\s\-–—_=*~·.,，。:;|/\\()\[\]{}<>]+", t):
                                return True
                        # 仅构成于数字/标点
                        if re.fullmatch(r"[0-9\s\-–—_=*~·.,，。:;|/\\()\[\]{}<>]+", t):
                            return True
                        core = re.findall(r"[A-Za-z0-9\u4e00-\u9fff]", t)
                        if not core:
                            return True
                        if len(t) <= 10 and all(c.isdigit() for c in core):
                            return True
                        ratio = len(core)/len(t)
                        if ratio < 0.18 and len(t) < 60:
                            return True
                        return False
                    filtered = []
                    for d in split_docs:
                        if _is_invalid(d.page_content or ''):
                            removed_invalid += 1
                            continue
                        filtered.append(d)
                    split_docs = filtered
                except Exception:
                    pass
            total_chars = sum(len(d.page_content or '') for d in split_docs)
            # 限制返回规模
            truncated = False
            if len(split_docs) > max_chunks:
                split_docs = split_docs[:max_chunks]
                truncated = True

            out_chunks: List[Dict[str, Any]] = []
            for idx, d in enumerate(split_docs):
                md = dict(getattr(d, 'metadata', {}) or {})
                out_chunks.append({
                    'index': idx,
                    'content': d.page_content,
                    'length': len(d.page_content or ''),
                    'start_index': md.get('start_index'),
                    'chunk_type': md.get('chunk_type'),
                    'h1': md.get('h1'),
                    'h2': md.get('h2'),
                    'h3': md.get('h3'),
                    'sheet': md.get('sheet'),
                    'row_index': md.get('row_index'),
                })

            return {
                'success': True,
                'chunk_size': chunk_size,
                'chunk_overlap': chunk_overlap,
                'semantic_refine': semantic_refine,
                'semantic_threshold': semantic_threshold,
                'chunk_count': len(out_chunks),
                'removed_invalid': removed_invalid,
                'original_chunk_count': original_chunk_count,
                'total_chars': total_chars,
                'truncated': truncated,
                'chunks': out_chunks,
            }
        except Exception as e:
            logger.warning(f"simulate_split failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def get_supported_file_types(self) -> List[str]:
        """
        获取支持的文件类型
        
        Returns:
            支持的文件扩展名列表
        """
        # 与上传白名单保持一致
        return ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.md', '.html', '.htm']
    
    def validate_file_type(self, filename: str) -> bool:
        """
        验证文件类型是否支持
        
        Args:
            filename: 文件名
            
        Returns:
            是否支持
        """
        file_ext = os.path.splitext(filename)[1].lower()
        return file_ext in self.get_supported_file_types()
