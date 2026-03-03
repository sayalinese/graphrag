from __future__ import annotations

import os
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    HTMLHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)
from typing import Optional, Tuple, List as _List
import math

try:
    import docx  # python-docx
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
except Exception:
    docx = None


def _is_markdown(ext: str) -> bool:
    return ext.lower() in {".md", ".markdown"}


def _is_html(ext: str) -> bool:
    return ext.lower() in {".html", ".htm"}


def split_markdown_with_headers(text: str) -> List[Document]:
    """
    按标题层级切分Markdown，同时保留标题上下文。
    每个块包含标题及其下的内容，避免标题被分散。
    """
    headers_to_split_on = [
        ("#", "h1"),
        ("##", "h2"),
        ("###", "h3"),
        ("####", "h4"),
        ("#####", "h5"),
        ("######", "h6"),
    ]
    splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
    docs = splitter.split_text(text)
    
    # 为每个文档块的内容前面加上它的所有标题上下文
    enriched_docs = []
    for doc in docs:
        md = doc.metadata or {}
        # 构建标题路径：h1 -> h1/h2 -> h1/h2/h3 等
        title_path = []
        for level in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            if level in md and md[level]:
                title_path.append(f"{'#' * (int(level[1]) - 1 + 1)} {md[level]}")
        
        # 如果有标题路径，在内容前加上
        if title_path:
            full_content = "\n".join(title_path) + "\n\n" + doc.page_content
        else:
            full_content = doc.page_content
        
        enriched_docs.append(Document(page_content=full_content, metadata=md))
    
    return enriched_docs


def split_html_with_headers(text: str) -> List[Document]:
    headers_to_split_on = [
        ("h1", "h1"),
        ("h2", "h2"),
        ("h3", "h3"),
        ("h4", "h4"),
        ("h5", "h5"),
        ("h6", "h6"),
    ]
    splitter = HTMLHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
    return splitter.split_text(text)


def smart_split_md_html(documents: List[Document], recursive_splitter: RecursiveCharacterTextSplitter | None = None) -> List[Document]:
    """
    对 Markdown/HTML 文本按标题层级优先切分；如单段过长，再用递归字符切分二次细化。

    documents: 输入 Document 列表（通常来自 TextLoader/HTML loader），需要包含 page_content；
    recursive_splitter: 可选的二次切分器（若未提供，将使用默认 300/100 配置）。
    """
    if not documents:
        return []

    # 若未提供，给一个默认递归切分器
    rec = recursive_splitter or RecursiveCharacterTextSplitter(
        chunk_size=600, chunk_overlap=120, length_function=len, add_start_index=True
    )

    out: List[Document] = []
    for doc in documents:
        text = (doc.page_content or "").strip()
        base_meta = dict(doc.metadata or {})
        source = base_meta.get("source") or ""
        ext = os.path.splitext(str(source))[-1].lower()

        try:
            sub_docs: List[Document]
            if _is_markdown(ext):
                sub_docs = split_markdown_with_headers(text)
            elif _is_html(ext):
                sub_docs = split_html_with_headers(text)
            else:
                # 非 md/html 直接跳过，由调用方走默认逻辑
                return []

            # 统一补充原始来源等元数据
            enriched: List[Document] = []
            for sd in sub_docs:
                md = dict(sd.metadata or {})
                # 保留原元数据
                for k, v in base_meta.items():
                    md.setdefault(k, v)
                    # 标准化 filename 字段
                    try:
                        import os as _os
                        if md.get('source'):
                            md['filename'] = _os.path.basename(str(md.get('source')))
                    except Exception:
                        pass
                enriched.append(Document(page_content=sd.page_content, metadata=md))

            # 对过长段落做二次细分（但要保留标题信息）
            for sd in enriched:
                content = sd.page_content
                if len(content) > 1200:
                    # 检查是否有标题前缀
                    lines = content.split('\n')
                    title_lines = []
                    body_start_idx = 0
                    
                    for i, line in enumerate(lines):
                        if line.strip().startswith('#'):
                            title_lines.append(line)
                        elif line.strip() == '':
                            # 空行可能在标题和内容之间
                            if title_lines and not any(c.strip() and not c.strip().startswith('#') for c in lines[:i]):
                                continue
                            body_start_idx = i + 1
                            break
                        else:
                            body_start_idx = i
                            break
                    
                    title_prefix = '\n'.join(title_lines)
                    body_content = '\n'.join(lines[body_start_idx:])
                    
                    # 对正文内容进行二次细分
                    if body_content.strip():
                        parts = rec.split_text(body_content)
                        for p in parts:
                            # 每个部分都加上标题前缀
                            final_content = (title_prefix + '\n\n' + p).strip() if title_prefix else p
                            out.append(Document(page_content=final_content, metadata=dict(sd.metadata)))
                    else:
                        # 如果只有标题，直接保留
                        out.append(Document(page_content=content, metadata=dict(sd.metadata)))
                else:
                    out.append(sd)
        except Exception:
            # 任意异常则回退（由调用方继续默认切分）
            return []

    return out


def _heading_level_from_style(style_name: Optional[str]) -> Optional[int]:
    if not style_name:
        return None
    name = str(style_name).strip().lower()
    # 常见：Heading 1/2/3，中文：标题 1/2/3
    for lvl in range(1, 7):
        if f"heading {lvl}" in name:
            return lvl
        if f"标题 {lvl}" in name or f"标题{lvl}" in name:
            return lvl
    return None


def _iter_docx_body(doc):
    """按顺序遍历 docx 文档体，依次产出 ('p', paragraph) 或 ('tbl', table)"""
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield 'p', docx.text.paragraph.Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield 'tbl', docx.table.Table(child, doc)


def smart_split_docx(documents: List[Document], recursive_splitter: RecursiveCharacterTextSplitter | None = None) -> List[Document]:
    """
    对 Word(.docx) 文档按 Heading 1/2/3 样式优先切分；对过长块用递归切分细化。
    需 metadata.source 指向原始 .docx 路径，且安装了 python-docx。
    返回 [] 表示放弃（由调用方回退到默认切分）。
    """
    if not documents:
        return []
    if docx is None:
        return []

    rec = recursive_splitter or RecursiveCharacterTextSplitter(
        chunk_size=300, chunk_overlap=100, length_function=len, add_start_index=True
    )

    out: List[Document] = []
    for d in documents:
        md0 = (d.metadata or {})
        # 若存在 processed_source（例如 .doc 转 .docx 的临时路径），优先用于结构化切分
        src = md0.get('processed_source') or md0.get('source')
        if not src or not str(src).lower().endswith('.docx'):
            # 仅处理 .docx；.doc 已在流程前被转成 .docx 再加载
            return []
        try:
            doc = docx.Document(str(src))
        except Exception:
            return []

        # 当前标题上下文
        h1 = h2 = h3 = None
        buf_lines: List[str] = []
        # 记录 start_index（在拼接的全文里模拟偏移）
        cursor = 0

        def flush_block():
            nonlocal buf_lines, cursor, h1, h2, h3
            if not buf_lines:
                return
            text = '\n'.join(buf_lines).strip()
            if not text:
                buf_lines = []
                return
            # 预览过滤仍希望以原始文件名匹配，若存在 original source 则保留；同时记录 processed_source
            md = {
                'source': str(md0.get('source') or src),
                'h1': h1,
                'h2': h2,
                'h3': h3,
                'start_index': cursor
            }
            # 透传 processed_source 以便调试
            if md0.get('processed_source'):
                md['processed_source'] = md0.get('processed_source')
            # 二次细分
            if len(text) > 600:
                parts = rec.split_text(text)
                local_offset = 0
                for p in parts:
                    out.append(Document(page_content=p, metadata={**md, 'start_index': cursor + local_offset, 'filename': os.path.basename(str(md.get('source'))) }))
                    local_offset += len(p)
            else:
                out.append(Document(page_content=text, metadata={**md, 'filename': os.path.basename(str(md.get('source'))) }))
            cursor += len(text) + 1  # 近似累计，+1 当作换行
            buf_lines = []

        for kind, item in _iter_docx_body(doc):
            if kind == 'p':
                text = item.text or ''
                style_name = getattr(getattr(item, 'style', None), 'name', None)
                lvl = _heading_level_from_style(style_name)
                if lvl is not None:
                    # 新标题，先冲刷已有正文块
                    flush_block()
                    if lvl == 1:
                        h1, h2, h3 = text.strip(), None, None
                    elif lvl == 2:
                        h2, h3 = text.strip(), None
                    elif lvl == 3:
                        h3 = text.strip()
                    else:
                        # 4+ 当作正文的分隔提示，直接加入缓冲
                        if text.strip():
                            buf_lines.append(text.strip())
                    continue
                # 普通段落
                if text.strip():
                    buf_lines.append(text.strip())
                else:
                    # 空段触发轻量分段
                    flush_block()
            elif kind == 'tbl':
                # 表格文本按行拼接
                rows_txt: List[str] = []
                try:
                    for r in item.rows:
                        cells = [c.text.strip() for c in r.cells]
                        rows_txt.append(' | '.join([c for c in cells if c]))
                except Exception:
                    pass
                if rows_txt:
                    buf_lines.append('\n'.join(rows_txt))

        # 末尾冲刷
        flush_block()

    # 如果没有产出，则回退
    if not out:
        return []
    return out


# PDF 智能切分（基于字体大小的标题启发式）
def smart_split_pdf(
    documents: List[Document],
    recursive_splitter: RecursiveCharacterTextSplitter | None = None,
) -> List[Document]:
    """
    对 PDF 文本尝试基于字体大小识别 h1/h2/h3，并按标题层级切分；对过长块再用递归字符切分。
    要求 metadata.source 指向原始 .pdf 路径；若 PyMuPDF 不可用或遇到异常，返回 [] 让调用方回退默认切分。
    """
    if not documents:
        return []
    # 找到一个 .pdf 源文件路径
    src = None
    base_meta = None
    for d in documents:
        md = d.metadata or {}
        s = md.get("source")
        if s and str(s).lower().endswith(".pdf"):
            src = str(s)
            base_meta = md
            break
    if not src:
        return []

    try:
        import fitz  # PyMuPDF
    except Exception:
        # 无法使用结构化信息，放弃（由调用方回退）
        return []

    # 递归切分器
    rec = recursive_splitter or RecursiveCharacterTextSplitter(
        chunk_size=600, chunk_overlap=120, length_function=len, add_start_index=True
    )

    # 第一遍统计字体大小分布
    sizes: _List[float] = []
    try:
        doc = fitz.open(src)
    except Exception:
        return []

    def _iter_lines():
        """遍历行，产出 (text, max_size) 按阅读顺序。"""
        for page in doc:
            try:
                data = page.get_text("dict")
            except Exception:
                continue
            for b in data.get("blocks", []) or []:
                for l in b.get("lines", []) or []:
                    line_text_parts = []
                    max_size = 0.0
                    for s in l.get("spans", []) or []:
                        t = (s.get("text") or "").strip()
                        if not t:
                            continue
                        line_text_parts.append(t)
                        try:
                            max_size = max(max_size, float(s.get("size") or 0.0))
                        except Exception:
                            pass
                    if line_text_parts:
                        text = "".join(line_text_parts).strip()
                        if text:
                            yield text, max_size

    # 收集尺寸
    for _t, _sz in _iter_lines():
        if _sz:
            sizes.append(_sz)

    if not sizes:
        # 字体尺寸不可用，放弃
        doc.close()
        return []

    sizes_sorted = sorted(sizes)
    median = sizes_sorted[len(sizes_sorted)//2]
    uniq = sorted(set(sizes), reverse=True)
    # 选择比中位数大的前 3 个尺寸作为候选标题层级
    head_sizes = [s for s in uniq if s >= (median + 0.8)]  # 0.8pt 以上视为标题
    h1 = head_sizes[0] if len(head_sizes) > 0 else None
    h2 = head_sizes[1] if len(head_sizes) > 1 else None
    h3 = head_sizes[2] if len(head_sizes) > 2 else None

    def _classify(sz: float) -> int:
        def near(a: Optional[float], b: float, tol: float = 0.9) -> bool:
            return a is not None and abs(a - b) <= tol
        if near(h1, sz):
            return 1
        if near(h2, sz):
            return 2
        if near(h3, sz):
            return 3
        return 0

    # 第二遍：按行生成块
    out: List[Document] = []
    h1_txt = h2_txt = h3_txt = None
    buf_lines: _List[str] = []
    cursor = 0

    def flush_block():
        nonlocal buf_lines, cursor, h1_txt, h2_txt, h3_txt
        if not buf_lines:
            return
        text = "\n".join(buf_lines).strip()
        if not text:
            buf_lines = []
            return
        md = dict(base_meta or {})
        # 标准化来源/文件名
        try:
            md['source'] = src
            md['filename'] = os.path.basename(src)
        except Exception:
            md['source'] = src
        md['h1'] = h1_txt
        md['h2'] = h2_txt
        md['h3'] = h3_txt
        md['start_index'] = cursor
        # 二次细分
        if len(text) > 600:
            parts = rec.split_text(text)
            local_offset = 0
            for p in parts:
                out.append(Document(page_content=p, metadata={**md, 'start_index': cursor + local_offset}))
                local_offset += len(p)
        else:
            out.append(Document(page_content=text, metadata=md))
        cursor += len(text) + 1
        buf_lines = []

    for line_text, line_size in _iter_lines():
        lvl = _classify(line_size)
        if lvl > 0:
            # 命中标题，先冲刷现有正文
            flush_block()
            if lvl == 1:
                h1_txt, h2_txt, h3_txt = line_text, None, None
            elif lvl == 2:
                h2_txt, h3_txt = line_text, None
            else:
                h3_txt = line_text
            continue
        # 正文行
        if not line_text.strip():
            flush_block()
        else:
            buf_lines.append(line_text)

    flush_block()
    doc.close()

    if not out:
        return []
    return out

# 语义边界检测

def _simple_sent_tokenize(text: str) -> _List[str]:
    """优先用 spaCy 切句；无可用模型时用规则兜底。"""
    # 尝试使用 spaCy
    try:
        import spacy  # type: ignore
        nlp = None
        # 优先中文模型
        for model in ("zh_core_web_sm", "zh_core_web_md", "zh_core_web_lg", "en_core_web_sm"):
            try:
                nlp = spacy.load(model)
                break
            except Exception:
                continue
        if nlp is not None:
            doc = nlp(text)
            sents = [s.text.strip() for s in doc.sents if s.text.strip()]
            if sents:
                return sents
    except Exception:
        pass
    # 规则兜底（中英混合）
    import re
    # 将常见句末标点作为断句依据，同时保留换行
    parts = re.split(r"(?<=[。！？!?；;。\.!?])\s+|\n+", text)
    sents = [p.strip() for p in parts if p and p.strip()]
    return sents if sents else [text]


def _cosine_sim(a: _List[float], b: _List[float]) -> float:
    if not a or not b:
        return 0.0
    if len(a) != len(b):
        # 粗糙兜底，长度不一致时取较小长度
        L = min(len(a), len(b))
        a = a[:L]
        b = b[:L]
    dot = sum(x*y for x, y in zip(a, b))
    na = math.sqrt(sum(x*x for x in a))
    nb = math.sqrt(sum(y*y for y in b))
    if na == 0.0 or nb == 0.0:
        return 0.0
    return dot / (na * nb)


def semantic_refine_with_embeddings(
    docs: _List[Document],
    embedding_service,  # EmbeddingService 或有 embeddings 属性的对象
    target_chunk_size: int = 800,
    overlap: int = 150,
    sim_threshold: float = 0.6,
    min_sentences: int = 1,
    max_sentences: Optional[int] = None,
) -> _List[Document]:
    """
    基于相邻句子相似度的边界检测，对较长文本做语义细化切分。

    - target_chunk_size/overlap：目标块大小与重叠（按字符估算，用于触发/限制合并）
    - sim_threshold：相邻句子相似度低于该阈值优先作为切点
    - min_sentences/max_sentences：限制每块句子数
    """
    if not docs:
        return []

    out: _List[Document] = []
    for d in docs:
        text = (d.page_content or "").strip()
        if not text:
            continue
        # 短文本直接保留
        if len(text) <= target_chunk_size:
            out.append(d)
            continue

        sents = _simple_sent_tokenize(text)
        if len(sents) <= 1:
            out.append(d)
            continue

        # 生成句向量（批量）
        try:
            emb = embedding_service.embeddings
            vectors = emb.embed_documents(sents)  # List[List[float]]
        except Exception:
            # 向量失败则保留原文
            out.append(d)
            continue

        # 计算相邻相似度
        sims = [_cosine_sim(vectors[i], vectors[i+1]) for i in range(len(vectors)-1)]

        # 合并成块：尽量在低相似度处断开，并不超过 target_chunk_size
        cur_buf: _List[str] = []
        cur_len = 0
        cur_sent_count = 0
        def flush_buf():
            nonlocal cur_buf, cur_len, cur_sent_count
            if not cur_buf:
                return
            content = "".join(cur_buf).strip()
            if content:
                # 继承元数据
                out.append(Document(page_content=content, metadata=dict(d.metadata or {})))
            cur_buf = []
            cur_len = 0
            cur_sent_count = 0

        for i, sent in enumerate(sents):
            s = sent
            s_len = len(s)
            # 若即将超出目标长度，尽量在上一个相似度低的位置断开
            should_cut = (cur_len + s_len) > max(target_chunk_size, 1)
            # 若相似度低于阈值，也考虑切分（对当前句之前断开）
            low_sim_here = (i > 0 and sims[i-1] < sim_threshold)

            cur_buf.append(s)
            cur_len += s_len
            cur_sent_count += 1

            # 满足最小句数后允许切分
            allow_cut = (cur_sent_count >= min_sentences)
            too_many = (max_sentences is not None and cur_sent_count >= max_sentences)

            if allow_cut and (too_many or should_cut or low_sim_here):
                flush_buf()
                # 追加重叠：简单取最后 overlap 字符为下一块的前缀
                if overlap > 0 and i < len(sents) - 1:
                    # 取上一块末尾 overlap 字符作为下一块开头（近似）
                    prefix = s[-overlap:]
                    if prefix:
                        cur_buf.append(prefix)
                        cur_len += len(prefix)

        # 冲刷最后缓冲
        flush_buf()

    return out


# Excel 智能切分：按 sheet + 表头 + 行拆分，每个块重复上下文，避免“无头”和半截
def smart_split_excel(
    documents: _List[Document],
    recursive_splitter: RecursiveCharacterTextSplitter | None = None,
    target_chunk_size: int = 800,
    overlap: int = 150,
) -> _List[Document]:
    """
    针对通过 pandas 生成的 Excel 文本（含 sheet/h3 元数据或 Markdown 表格），
    按行拆成块，并将行的关键信息（除最长列外）作为前缀上下文重复到每个子块，
    防止“无头”与中途截断造成语义丢失。

    返回 [] 表示放弃，由调用方回退默认切分。
    """
    if not documents:
        return []

    out: _List[Document] = []

    def _is_candidate_excel_doc(d: Document) -> bool:
        md = d.metadata or {}
        # 我们在加载时会写入 sheet/h2/h3；或者文本里有 Markdown 表格痕迹
        txt = (d.page_content or "").lstrip()
        return bool(md.get('sheet') or md.get('h2') or (txt.startswith('Sheet:') and '| ' in txt))

    def _split_md_table(text: str) -> Tuple[_List[str], _List[_List[str]]]:
        """从 Markdown 表格提取 header 与 rows。若失败返回空。"""
        lines = [ln.strip() for ln in text.splitlines()]
        # 寻找第一条以 | 开头且包含分隔的行作为 header
        header_idx = None
        for i, ln in enumerate(lines):
            if ln.startswith('|') and '|' in ln.strip('|'):
                header_idx = i
                break
        if header_idx is None or header_idx + 1 >= len(lines):
            return [], []
        header_line = lines[header_idx]
        # 下一行应为 --- 分隔
        divider_line = lines[header_idx + 1] if header_idx + 1 < len(lines) else ''
        if '---' not in divider_line:
            return [], []
        def split_bar(ln: str) -> _List[str]:
            parts = [p.strip() for p in ln.strip().strip('|').split('|')]
            return parts
        headers = split_bar(header_line)
        rows: _List[_List[str]] = []
        for ln in lines[header_idx + 2:]:
            if not ln.startswith('|'):
                # 表格结束
                break
            vals = split_bar(ln)
            # 对齐列数
            if len(vals) < len(headers):
                vals += [''] * (len(headers) - len(vals))
            elif len(vals) > len(headers):
                vals = vals[:len(headers)]
            rows.append(vals)
        return headers, rows

    rec = recursive_splitter or RecursiveCharacterTextSplitter(
        chunk_size=target_chunk_size,
        chunk_overlap=overlap,
        length_function=len,
        add_start_index=True,
    )

    for d in documents:
        if not _is_candidate_excel_doc(d):
            # 非 excel 风格文本，放弃整个策略（由调用方回退）
            return []

        base_md = dict(d.metadata or {})
        sheet_name = base_md.get('sheet') or base_md.get('h2')
        text = d.page_content or ''
        # 查找 Markdown 表格
        headers, rows = _split_md_table(text)
        if not headers or not rows:
            # 无法识别表格，放弃
            return []

        # 识别最长列，用来当作主体（例如“回答”列）
        col_lengths = [0] * len(headers)
        for r in rows:
            for i, v in enumerate(r):
                col_lengths[i] = max(col_lengths[i], len(v or ''))
        long_col_idx = max(range(len(headers)), key=lambda i: col_lengths[i]) if headers else 0

        # 标准化列名，定位问答列；过滤无意义列（如 Unnamed: 0）
        def _norm(s: str) -> str:
            return (s or '').strip()

        def _is_unnamed(s: str) -> bool:
            ss = (s or '').strip().lower()
            return ss.startswith('unnamed') or ss in {'', 'index'}

        norm_headers = [_norm(h) for h in headers]
        # 常见中文问答列名匹配
        def _find_idx(candidates: _List[str]) -> Optional[int]:
            for c in candidates:
                for i, h in enumerate(norm_headers):
                    if c in h:
                        return i
            return None

        q_idx = _find_idx(['提问', '问题', '问'])
        a_idx = _find_idx(['回答', '答案', '答'])
        # 用于展示的列名映射（更友好的标签）
        def _display_key(h: str) -> str:
            hn = _norm(h)
            if '提问' in hn or hn == '问' or '问题' in hn:
                return '问题'
            if '回答' in hn or hn == '答' or '答案' in hn:
                return '回答'
            return hn

        # 计算粗略 start_index：在原文本中的位置（基于行长度累加）
        cursor = 0
        for ln in text.splitlines():
            # 找到 header 行位置，cursor 停留到 header 后一行（divider）之后
            if ln.strip().startswith('|') and '|' in ln.strip().strip('|'):
                # header + divider 两行
                cursor += len(ln) + 1
                # 加 divider 下一行
                # 安全处理：尝试取下一行长度
                break
            cursor += len(ln) + 1
        # 再加上 divider 行长度
        # 简化：不再精确定位，后续每行按累加即可

        for row_idx, vals in enumerate(rows, start=1):
            # 组装上下文前缀（除最长列、且若识别到问答列也排除“问题”列，避免重复）
            kv_pairs = []
            for i, h in enumerate(headers):
                if i == long_col_idx:
                    continue
                v = vals[i] if i < len(vals) else ''
                if v:
                    # 过滤无意义列（如 Unnamed: *）
                    if _is_unnamed(h):
                        continue
                    # 如果稍后会在前缀追加“问题: ...”，这里避免再次包含
                    if q_idx is not None and i == q_idx:
                        continue
                    kv_pairs.append(f"{_display_key(h)}: {v}")
            # 前缀不再包含 Sheet 行，避免干扰语义；仅保留有意义的键值
            prefix_lines = []
            if kv_pairs:
                prefix_lines.extend(kv_pairs)
            prefix = "\n".join(prefix_lines).strip()

            # 优先以 问题/回答 结构输出，更利于中文检索
            if q_idx is not None and a_idx is not None and q_idx < len(vals) and a_idx < len(vals):
                question = (vals[q_idx] or '').strip()
                answer = (vals[a_idx] or '').strip()
                body_label = '回答'
                body_text = answer
                # 构建固定前缀：类别等上下文 + 问题
                qa_prefix_lines = []
                if prefix:
                    qa_prefix_lines.append(prefix)
                if question:
                    qa_prefix_lines.append(f"问题: {question}")
                qa_prefix = "\n\n".join([p for p in qa_prefix_lines if p])
            else:
                body_label = headers[long_col_idx] if long_col_idx < len(headers) else ''
                body_text = vals[long_col_idx] if long_col_idx < len(vals) else ''
                body_text = (body_text or '').strip()
                qa_prefix = prefix

            def emit_chunk(seg_text: str, local_offset: int = 0):
                # 内容采用「前缀上下文 + 主体」；若识别为问答结构，则主体标签固定为“回答”，前缀包含“问题”
                content_prefix = qa_prefix
                content = (content_prefix + ("\n\n" if content_prefix and seg_text else "") + (f"{body_label}: " if body_label and seg_text else '') + seg_text).strip()
                md = {**base_md}
                md.setdefault('h2', sheet_name)
                md.setdefault('h3', " | ".join(headers) if headers else None)
                md['start_index'] = cursor + local_offset
                try:
                    import os as _os
                    if md.get('source'):
                        md['filename'] = _os.path.basename(str(md.get('source')))
                except Exception:
                    pass
                out.append(Document(page_content=content, metadata=md))

            if not body_text:
                # 纯上下文一块
                emit_chunk("")
                # 粗略推进 cursor：当前行长度
                cursor += sum(len(str(v)) for v in vals) + len(vals)  # 近似
                continue

            # 若主体很长，则分段，同时重复前缀，避免“无头”
            max_body = max(1, target_chunk_size - max(80, len(qa_prefix)))
            if len(body_text) <= max_body:
                emit_chunk(body_text, local_offset=0)
            else:
                # 简单按字符窗口切分，并加入 overlap
                start = 0
                while start < len(body_text):
                    end = min(len(body_text), start + max_body)
                    seg = body_text[start:end]
                    emit_chunk(seg, local_offset=start)
                    if end >= len(body_text):
                        break
                    # 处理重叠，确保始终向前推进
                    start = max(start + 1, end - overlap)

            # 推进 cursor：按近似行长度累加
            cursor += sum(len(str(v)) for v in vals) + len(vals)

    return out
